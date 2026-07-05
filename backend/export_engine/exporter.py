import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json
import os
import zipfile
from typing import Dict, Any, Optional
import logging
from datetime import datetime

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class ResultExporter:
    def __init__(self):
        self.export_dir = "exports"
        os.makedirs(self.export_dir, exist_ok=True)
    
    def export_all(self, empirical_results: Dict[str, Any], paper_content: str, project_id: str) -> str:
        """
        导出所有结果
        """
        try:
            # 创建项目导出目录
            project_dir = os.path.join(self.export_dir, f"project_{project_id}_{datetime.now().strftime('%Y%m%d%H%M%S')}")
            os.makedirs(project_dir, exist_ok=True)
            
            # 导出原始数据结果
            self._export_raw_results(empirical_results, project_dir)
            
            # 导出Word论文
            word_path = self._export_word_paper(empirical_results, paper_content, project_dir)
            
            # 导出图表
            self._export_charts(empirical_results, project_dir)
            
            # 打包所有文件
            zip_path = self._create_zip(project_dir, project_id)
            
            logger.info(f"所有结果导出完成，路径: {zip_path}")
            return zip_path
            
        except Exception as e:
            logger.error(f"导出失败: {str(e)}")
            raise
    
    def _export_raw_results(self, results: Dict[str, Any], output_dir: str):
        """
        导出原始计算结果
        """
        # 导出JSON格式结果
        json_path = os.path.join(output_dir, "empirical_results.json")
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(results, f, ensure_ascii=False, indent=2)
        
        # 导出描述性统计
        if "descriptive" in results:
            desc_df = pd.DataFrame.from_dict(results["descriptive"]["variables"]).T
            desc_path = os.path.join(output_dir, "descriptive_statistics.csv")
            desc_df.to_csv(desc_path, encoding='utf-8-sig')
        
        # 导出相关性分析
        if "correlation" in results:
            corr_df = pd.DataFrame.from_dict(results["correlation"]["pearson"])
            corr_path = os.path.join(output_dir, "correlation_matrix.csv")
            corr_df.to_csv(corr_path, encoding='utf-8-sig')
            
            pvalue_df = pd.DataFrame.from_dict(results["correlation"]["p_values"])
            pvalue_path = os.path.join(output_dir, "correlation_pvalues.csv")
            pvalue_df.to_csv(pvalue_path, encoding='utf-8-sig')
        
        # 导出回归结果
        if "regression" in results:
            reg_results = results["regression"]
            for model_name, model_result in reg_results.items():
                if isinstance(model_result, dict) and "coefficients" in model_result:
                    coef_df = pd.DataFrame.from_dict(model_result["coefficients"]).T
                    coef_path = os.path.join(output_dir, f"regression_coefficients_{model_name}.csv")
                    coef_df.to_csv(coef_path, encoding='utf-8-sig')
                    
                    # 导出回归摘要
                    summary_path = os.path.join(output_dir, f"regression_summary_{model_name}.txt")
                    with open(summary_path, 'w', encoding='utf-8') as f:
                        if "summary" in model_result:
                            f.write(model_result["summary"])
    
    def _export_word_paper(self, results: Dict[str, Any], paper_content: str, output_dir: str) -> str:
        """
        导出Word格式论文
        """
        doc = Document()
        
        # 设置文档样式
        self._set_document_style(doc)
        
        # 添加标题
        doc.add_heading('实证研究论文', level=1)
        
        # 添加作者和日期
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run = p.add_run(f'生成日期: {datetime.now().strftime("%Y年%m月%d日")}\n')
        run.font.size = Pt(10)
        
        # 添加目录
        doc.add_heading('目录', level=2)
        doc.add_paragraph('1. 数据处理说明', style='List Number')
        doc.add_paragraph('2. 描述性统计分析', style='List Number')
        doc.add_paragraph('3. 相关性分析', style='List Number')
        doc.add_paragraph('4. 回归结果分析', style='List Number')
        doc.add_paragraph('5. 稳健性检验', style='List Number')
        doc.add_paragraph('6. 研究结论与启示', style='List Number')
        
        # 添加AI生成的论文内容
        doc.add_page_break()
        doc.add_heading('实证研究内容', level=2)
        
        # 分割AI生成的内容并添加到文档
        content_sections = paper_content.split('\n\n')
        for section in content_sections:
            if section.strip():
                p = doc.add_paragraph(section.strip())
                p.paragraph_format.line_spacing = 1.5
        
        # 添加附录：原始结果表格
        doc.add_page_break()
        doc.add_heading('附录：实证结果表格', level=2)
        
        # 添加描述性统计表格
        if "descriptive" in results:
            doc.add_heading('1. 描述性统计', level=3)
            self._add_descriptive_table(doc, results["descriptive"])
        
        # 添加相关性分析表格
        if "correlation" in results:
            doc.add_heading('2. 相关性分析', level=3)
            self._add_correlation_table(doc, results["correlation"])
        
        # 添加回归结果表格
        if "regression" in results:
            doc.add_heading('3. 回归结果', level=3)
            self._add_regression_tables(doc, results["regression"])
        
        # 保存文档
        word_path = os.path.join(output_dir, "empirical_paper.docx")
        doc.save(word_path)
        
        return word_path
    
    def _set_document_style(self, doc):
        """
        设置文档样式
        """
        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
        font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 设置标题样式
        for level in [1, 2, 3, 4, 5, 6]:
            style = doc.styles[f'Heading {level}']
            font = style.font
            font.name = '黑体'
            font.bold = True
            font.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    def _add_descriptive_table(self, doc, descriptive_results):
        """
        添加描述性统计表格
        """
        desc_df = pd.DataFrame.from_dict(descriptive_results["variables"]).T
        
        table = doc.add_table(rows=1, cols=len(desc_df.columns) + 1)
        table.style = 'Table Grid'
        
        # 设置表格标题
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '变量'
        for i, col in enumerate(desc_df.columns):
            hdr_cells[i+1].text = str(col)
        
        # 添加数据行
        for var_name, row in desc_df.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = var_name
            for i, val in enumerate(row):
                if isinstance(val, (int, float)):
                    row_cells[i+1].text = f'{val:.4f}' if isinstance(val, float) else str(val)
                else:
                    row_cells[i+1].text = str(val)
    
    def _add_correlation_table(self, doc, correlation_results):
        """
        添加相关性分析表格
        """
        corr_df = pd.DataFrame.from_dict(correlation_results["pearson"])
        
        table = doc.add_table(rows=1, cols=len(corr_df.columns) + 1)
        table.style = 'Table Grid'
        
        # 设置表格标题
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '变量'
        for i, col in enumerate(corr_df.columns):
            hdr_cells[i+1].text = str(col)
        
        # 添加数据行
        for var_name, row in corr_df.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = var_name
            for i, val in enumerate(row):
                row_cells[i+1].text = f'{val:.4f}'
                # 添加显著性标记
                p_val = correlation_results["p_values"][var_name][corr_df.columns[i]]
                if p_val < 0.01:
                    row_cells[i+1].text += '***'
                elif p_val < 0.05:
                    row_cells[i+1].text += '**'
                elif p_val < 0.1:
                    row_cells[i+1].text += '*'
    
    def _add_regression_tables(self, doc, regression_results):
        """
        添加回归结果表格
        """
        for model_name, model_result in regression_results.items():
            if isinstance(model_result, dict) and "coefficients" in model_result:
                doc.add_heading(f'{model_result["model_name"]} 回归结果', level=4)
                
                coef_df = pd.DataFrame.from_dict(model_result["coefficients"]).T
                
                table = doc.add_table(rows=1, cols=len(coef_df.columns) + 1)
                table.style = 'Table Grid'
                
                # 设置表格标题
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = '变量'
                for i, col in enumerate(coef_df.columns):
                    hdr_cells[i+1].text = str(col)
                
                # 添加数据行
                for var_name, row in coef_df.iterrows():
                    row_cells = table.add_row().cells
                    row_cells[0].text = var_name
                    for i, val in enumerate(row):
                        if isinstance(val, (int, float)):
                            row_cells[i+1].text = f'{val:.4f}' if isinstance(val, float) else str(val)
                        else:
                            row_cells[i+1].text = str(val)
                
                # 添加模型统计量
                p = doc.add_paragraph()
                run = p.add_run('模型统计量: ')
                run.bold = True
                
                stats = model_result["model_stats"]
                stats_text = []
                if stats["r_squared"] is not None:
                    stats_text.append(f'R² = {stats["r_squared"]:.4f}')
                if stats["adj_r_squared"] is not None:
                    stats_text.append(f'调整R² = {stats["adj_r_squared"]:.4f}')
                if stats["f_statistic"] is not None:
                    stats_text.append(f'F值 = {stats["f_statistic"]:.4f}')
                if stats["n_obs"] is not None:
                    stats_text.append(f'样本量 = {int(stats["n_obs"])}')
                
                p.add_run(' '.join(stats_text))
    
    def _export_charts(self, results: Dict[str, Any], output_dir: str):
        """
        导出可视化图表
        """
        # 这里可以集成ECharts生成图表
        # 目前先保存图表数据，供前端可视化使用
        
        chart_data = {}
        
        # 描述性统计图表数据
        if "descriptive" in results:
            chart_data["descriptive"] = results["descriptive"]
        
        # 相关性分析图表数据
        if "correlation" in results:
            chart_data["correlation"] = results["correlation"]
        
        # 回归结果图表数据
        if "regression" in results:
            chart_data["regression"] = results["regression"]
        
        # 保存图表数据
        chart_path = os.path.join(output_dir, "chart_data.json")
        with open(chart_path, 'w', encoding='utf-8') as f:
            json.dump(chart_data, f, ensure_ascii=False, indent=2)
    
    def _create_zip(self, source_dir: str, project_id: str) -> str:
        """
        打包所有导出文件
        """
        zip_filename = f"empirical_results_{project_id}_{datetime.now().strftime('%Y%m%d%H%M%S')}.zip"
        zip_path = os.path.join(self.export_dir, zip_filename)
        
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            # 遍历目录中的所有文件
            for root, dirs, files in os.walk(source_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, source_dir)
                    zipf.write(file_path, arcname)
        
        return zip_path
